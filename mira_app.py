"""
Mira v1.5 — Real Estate Transaction Co-Pilot Backend
Major additions for Hosea's workflow:
- Real PDF contract autopopulation (Virginia REIN forms)
- Realist MLS data extraction and integration
- Transaction deadline tracking foundation
- Enhanced lead status pipeline
"""

import os
import json
import datetime
from typing import Optional, Dict, Any
from fastapi import FastAPI, Request, BackgroundTasks, Body, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.ext.asyncio import create_async_engine
from sqlalchemy import text
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO

# Single FastAPI app instance
app = FastAPI()

# Templates setup
templates = Jinja2Templates(directory="templates")

# Database helper function
def get_database_url():
    DATABASE_URL = os.getenv("DATABASE_URL")
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql+asyncpg://", 1)
    elif DATABASE_URL.startswith("postgresql://"):
        DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgresql+asyncpg://", 1)
    return DATABASE_URL

# ===== FIXED: Working PDF overlay function =====
def autopopulate_purchase_agreement(input_pdf_path, output_pdf_path, fields: Dict[str, str]):
    """
    Page 1 overlay for Virginia REIN Purchase Agreement
    Writes Buyer, Property Address, Price, MLS# into correct blanks.
    """
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    # Set small, contract-friendly font
    can.setFont("Helvetica", 9)

    # === Page 1 Field Mappings (coordinates in points) ===
    # Note: (0,0) is bottom-left, Letter = 612x792

    # Buyer Name (top left, line 1)
    can.drawString(90, 735, fields.get("buyer_name", ""))

    # Property Address line
    can.drawString(140, 695, fields.get("property_address", ""))

    # Purchase Price box (bottom left ~ line 30)
    can.drawString(100, 560, fields.get("price", ""))

    # MLS # small box on right
    can.drawString(500, 695, fields.get("mls", ""))

    can.save()
    packet.seek(0)

    # ===== Merge with Template =====
    overlay_pdf = PdfReader(packet)
    base_pdf = PdfReader(open(input_pdf_path, "rb"))
    writer = PdfWriter()

    # Merge page 1 with overlay
    base_page = base_pdf.pages[0]
    base_page.merge_page(overlay_pdf.pages[0])
    writer.add_page(base_page)

    # Copy rest of base unchanged
    for i in range(1, len(base_pdf.pages)):
        writer.add_page(base_pdf.pages[i])

    with open(output_pdf_path, "wb") as out_f:
        writer.write(out_f)

# ===== NEW: REALIST PDF PARSER =====
def extract_realist_data(pdf_file_path: str) -> Dict[str, Any]:
    """
    Extract property data from Realist MLS PDF
    Includes Tax ID, legal desc, subdivision, zoning, owners, etc.
    """
    extracted_data = {
        "property_address": "",
        "listing_price": "",
        "mls_number": "",
        "square_footage": "",
        "lot_size": "",
        "year_built": "",
        "bedrooms": "",
        "bathrooms": "",
        "property_type": "",
        "tax_id": "",
        "legal_description": "",
        "subdivision": "",
        "zoning": "",
        "assessed_value": "",
        "owner_of_record": "",
        "listing_agent": "",
        "listing_office": ""
    }
    
    try:
        with pdfplumber.open(pdf_file_path) as pdf:
            text_content = ""
            for page_num in range(min(3, len(pdf.pages))):
                text_content += pdf.pages[page_num].extract_text() or ""
            
            lines = text_content.split('\n')
            for line in lines:
                line_lower = line.lower().strip()
                
                if 'address' in line_lower and any(char.isdigit() for char in line):
                    extracted_data["property_address"] = line.strip()
                if '$' in line and ('price' in line_lower or 'list' in line_lower):
                    extracted_data["listing_price"] = line.strip()
                if 'mls' in line_lower and any(char.isdigit() for char in line):
                    extracted_data["mls_number"] = line.strip()
                if 'sq' in line_lower and 'ft' in line_lower:
                    extracted_data["square_footage"] = line.strip()
                if 'lot size' in line_lower:
                    extracted_data["lot_size"] = line.strip()
                if 'year built' in line_lower:
                    extracted_data["year_built"] = line.strip()
                if 'bed' in line_lower:
                    extracted_data["bedrooms"] = line.strip()
                if 'bath' in line_lower:
                    extracted_data["bathrooms"] = line.strip()
                if 'type' in line_lower:
                    extracted_data["property_type"] = line.strip()
                if 'tax id' in line_lower or 'parcel' in line_lower:
                    extracted_data["tax_id"] = line.strip()
                if 'legal' in line_lower:
                    extracted_data["legal_description"] = line.strip()
                if 'subdivision' in line_lower:
                    extracted_data["subdivision"] = line.strip()
                if 'zoning' in line_lower:
                    extracted_data["zoning"] = line.strip()
                if 'assessed' in line_lower or 'assessment' in line_lower:
                    extracted_data["assessed_value"] = line.strip()
                if 'owner' in line_lower:
                    extracted_data["owner_of_record"] = line.strip()
    
    except Exception as e:
        print(f"Error extracting Realist data: {e}")
    
    return extracted_data

# ===== NEW: DEADLINE TRACKING SYSTEM =====
def calculate_transaction_deadlines(contract_date: datetime.date, contract_type: str = "purchase") -> Dict[str, datetime.date]:
    """
    NEW FUNCTION: Calculate critical deadlines based on contract execution date
    Virginia standard timelines (can be customized per contract)
    """
    deadlines = {}
    
    if contract_type == "purchase":
        # Standard Virginia purchase agreement timelines
        deadlines["inspection_period"] = contract_date + datetime.timedelta(days=10)
        deadlines["financing_contingency"] = contract_date + datetime.timedelta(days=21)
        deadlines["appraisal_contingency"] = contract_date + datetime.timedelta(days=21)
        deadlines["settlement_date"] = contract_date + datetime.timedelta(days=30)
        deadlines["title_commitment"] = contract_date + datetime.timedelta(days=15)
    
    return deadlines

# ===== FIXED: CONTRACT GENERATION =====
def generate_real_contract(lead_dict: Dict, realist_data: Dict = None) -> str:
    """
    FIXED: Generate actual Virginia REIN contract with proper overlay merge
    Integrates lead data + Realist property data
    """
    template_path = os.path.join("templates", "contracts", "Standard_Purchase_Agreement.pdf")
    print(f"DEBUG: Looking for template at {template_path}")
    
    if not os.path.exists(template_path):
        print("DEBUG: Template not found, falling back to demo contract")
        # Fallback to demo generation if template not found
        return generate_demo_contract(lead_dict)
    
    os.makedirs("generated_contracts", exist_ok=True)
    output_path = f"generated_contracts/PA_filled_{lead_dict.get('id', 'unknown')}.pdf"
    
    if realist_data is None:
        realist_data = {}
    
    # Prepare fields for overlay
    fields = {
        "buyer_name": lead_dict.get("name", ""),
        "email": lead_dict.get("email", ""),
        "property_address": realist_data.get("property_address", ""),
        "price": realist_data.get("listing_price", ""),
        "mls": realist_data.get("mls_number", "")
    }
    
    # Use the fixed overlay function
    autopopulate_purchase_agreement(template_path, output_path, fields)
    return output_path

# Original demo contract function (kept as fallback)
def generate_demo_contract(lead_dict):
    """Generate a demo contract for the given lead"""
    from docx import Document
    
    os.makedirs("generated_contracts", exist_ok=True)
    
    doc = Document()
    doc.add_heading(f"Service Agreement - {lead_dict['name']}", 0)
    doc.add_paragraph(f"Client: {lead_dict['name']}")
    doc.add_paragraph(f"Email: {lead_dict['email']}")
    doc.add_paragraph(f"Service: {lead_dict['service']}")
    doc.add_paragraph("This is a demo contract generated by Mira AI Real Estate Co-Pilot.")
    
    file_path = f"generated_contracts/demo_contract_{lead_dict['id']}.docx"
    doc.save(file_path)
    return file_path

# Root endpoint
@app.get("/")
async def root():
    return {"status": "Mira v1.5 backend is alive 🚀 - Now with real contract autopopulation"}

# ===== ENHANCED: DASHBOARD WITH NEW STATUS PIPELINE =====
@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.connect() as conn:
        # NEW: Enhanced query to include Realist data and deadlines
        result = await conn.execute(text("""
            SELECT id, name, email, service, status, realist_data, created_at 
            FROM leads ORDER BY created_at DESC
        """))
        raw_leads = result.fetchall()

    await engine.dispose()

    # NEW: Enhanced status pipeline for real transactions
    leads_by_status = {
        "🆕 New Leads": [],
        "📋 Realist Data Added": [],
        "📄 Contract Drafted": [],
        "👀 Awaiting Agent Review": [],
        "✍️ DocuSign Ready": [],
        "⏰ Pending Signatures": [],
        "✅ Completed": [],
        "❓ Needs Attention": []
    }

    # Enhanced status mapping
    status_map = {
        "new": "🆕 New Leads",
        "realist_added": "📋 Realist Data Added",
        "contract_drafted": "📄 Contract Drafted",
        "awaiting_review": "👀 Awaiting Agent Review",
        "docusign_ready": "✍️ DocuSign Ready",
        "pending_signatures": "⏰ Pending Signatures",
        "completed": "✅ Completed"
    }

    for lead in raw_leads:
        lead_dict = {
            "id": lead[0],
            "name": lead[1],
            "email": lead[2],
            "service": lead[3],
            "status": lead[4],
            "realist_data": lead[5],  # NEW: Realist data column
            "created_at": lead[6]
        }

        normalized = (lead_dict["status"] or "").strip().lower()
        status_key = status_map.get(normalized, "❓ Needs Attention")
        leads_by_status[status_key].append(lead_dict)

    return templates.TemplateResponse(
        "dashboard.html",
        {"request": request, "leads": leads_by_status}
    )

# Original Tally webhook (unchanged)
@app.post("/tally_webhook")
async def tally_webhook(payload: dict = Body(...)):
    print("📩 Incoming Tally Webhook Payload:", payload)

    name, email, service = "Unknown", "unknown@example.com", "General Inquiry"
    
    for ans in payload.get("data", {}).get("fields", []):
        label = ans.get("label", "").lower()
        if "full legal name" in label:
            name = ans.get("value") or name
        elif label == "email" and ans.get("value"):
            email = ans.get("value")
        elif "how can mira help you today?" in label:
            choice_ids = ans.get("value", [])
            options = {opt["id"]: opt["text"] for opt in ans.get("options", [])}
            if choice_ids:
                service = options.get(choice_ids[0], service)

    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.begin() as conn:
        await conn.execute(
            text("""
                INSERT INTO leads (name, email, service, status, raw_data)
                VALUES (:name, :email, :service, :status, :raw_data)
            """),
            {
                "name": name,
                "email": email,
                "service": service,
                "status": "new",
                "raw_data": json.dumps(payload)
            }
        )

    await engine.dispose()
    return {"success": True, "inserted": {"name": name, "email": email, "service": service}}

# ===== NEW: REALIST PDF UPLOAD ENDPOINT =====
@app.post("/upload_realist/{lead_id}")
async def upload_realist_data(lead_id: int, file: UploadFile = File(...)):
    """
    NEW ENDPOINT: Upload and parse Realist MLS PDF for property data
    Stores extracted data in database for contract autopopulation
    """
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Save uploaded file temporarily
    temp_path = f"temp_realist_{lead_id}.pdf"
    with open(temp_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Extract property data
    realist_data = extract_realist_data(temp_path)
    
    # Store in database
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)
    
    async with engine.begin() as conn:
        await conn.execute(
            text("""
                UPDATE leads 
                SET realist_data = :realist_data, status = :status 
                WHERE id = :id
            """),
            {
                "realist_data": json.dumps(realist_data),
                "status": "realist_added",
                "id": lead_id
            }
        )
    
    await engine.dispose()
    
    # Clean up temp file
    os.remove(temp_path)
    
    return {
        "success": True,
        "lead_id": lead_id,
        "extracted_data": realist_data,
        "message": "Realist data uploaded and parsed successfully"
    }

# ===== ENHANCED: REAL CONTRACT GENERATION =====
@app.post("/generate_contract/{lead_id}")
async def generate_contract(lead_id: int):
    """
    ENHANCED: Generate real Virginia REIN contract with autopopulated data
    Combines lead intake + Realist property data
    """
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.connect() as conn:
        result = await conn.execute(
            text("SELECT id, name, email, service, status, realist_data FROM leads WHERE id = :id"),
            {"id": lead_id}
        )
        row = result.fetchone()

    await engine.dispose()

    if not row:
        raise HTTPException(status_code=404, detail=f"Lead with id {lead_id} not found")

    # Convert SQLAlchemy row → dict
    lead_dict = {
        "id": row[0],
        "name": row[1],
        "email": row[2],
        "service": row[3],
        "status": row[4],
        "realist_data": row[5]
    }

    # Parse Realist data if available
    realist_data = {}
    if lead_dict["realist_data"]:
        try:
            realist_data = json.loads(lead_dict["realist_data"])
        except:
            realist_data = {}

    # NEW: Generate real contract instead of demo
    contract_path = generate_real_contract(lead_dict, realist_data)
    
    # NEW: Calculate transaction deadlines
    contract_date = datetime.date.today()
    deadlines = calculate_transaction_deadlines(contract_date)
    
    # NEW: Update lead status to contract_drafted
    engine = create_async_engine(DATABASE_URL, echo=False)
    async with engine.begin() as conn:
        await conn.execute(
            text("UPDATE leads SET status = :status WHERE id = :id"),
            {"status": "contract_drafted", "id": lead_id}
        )
    await engine.dispose()

    return {
        "success": True,
        "contract_created": contract_path,
        "lead": lead_dict,
        "realist_data": realist_data,
        "deadlines": deadlines,  # NEW: Include calculated deadlines
        "message": "Real estate contract generated and ready for agent review"
    }

# ===== NEW: GET LEAD DETAILS WITH REALIST DATA =====
@app.get("/lead/{lead_id}")
async def get_lead_details(lead_id: int):
    """
    NEW ENDPOINT: Get full lead details including Realist data and deadlines
    For agent review before contract finalization
    """
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.connect() as conn:
        result = await conn.execute(
            text("SELECT * FROM leads WHERE id = :id"),
            {"id": lead_id}
        )
        row = result.fetchone()

    await engine.dispose()

    if not row:
        raise HTTPException(status_code=404, detail=f"Lead {lead_id} not found")

    # Parse stored data
    realist_data = {}
    if row[5]:  # realist_data column
        try:
            realist_data = json.loads(row[5])
        except:
            pass

    return {
        "lead_id": row[0],
        "name": row[1],
        "email": row[2],
        "service": row[3],
        "status": row[4],
        "realist_data": realist_data,
        "created_at": row[6],
        "raw_intake_data": json.loads(row[7]) if row[7] else {}
    }

# Enhanced download endpoint
@app.get("/download_contract/{lead_id}")
async def download_contract(lead_id: int):
    pdf_path = f"generated_contracts/PA_filled_{lead_id}.pdf"
    if os.path.exists(pdf_path):
        return FileResponse(
            path=pdf_path,
            filename=f"purchase_agreement_{lead_id}.pdf",
            media_type="application/pdf"
        )

    raise HTTPException(status_code=404, detail=f"No PDF contract found for lead {lead_id}")

@app.post("/generate_and_download/{lead_id}")
async def generate_and_download_contract(lead_id: int):
    """
    NEW ENDPOINT: Generate a fresh Virginia REIN contract for this lead
    then immediately return the new PDF for download.
    """
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    # 1. Grab lead data
    async with engine.connect() as conn:
        result = await conn.execute(
            text("SELECT id, name, email, service, status, realist_data FROM leads WHERE id = :id"),
            {"id": lead_id}
        )
        row = result.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail=f"Lead with id {lead_id} not found")

    lead_dict = {
        "id": row[0],
        "name": row[1],
        "email": row[2],
        "service": row[3],
        "status": row[4],
        "realist_data": row[5]
    }

    realist_data = {}
    if lead_dict["realist_data"]:
        try:
            realist_data = json.loads(lead_dict["realist_data"])
        except:
            realist_data = {}

    # 2. Generate a fresh PDF contract
    contract_path = generate_real_contract(lead_dict, realist_data)

    # 3. Update DB status → contract_drafted
    async with engine.begin() as conn:
        await conn.execute(
            text("UPDATE leads SET status = :status WHERE id = :id"),
            {"status": "contract_drafted", "id": lead_id}
        )
    await engine.dispose()

    # 4. Immediately return the new PDF as a download
    if os.path.exists(contract_path):
        return FileResponse(
            path=contract_path,
            filename=f"purchase_agreement_{lead_id}.pdf",
            media_type="application/pdf"
        )
    else:
        raise HTTPException(status_code=500, detail="PDF generation failed")

# ===== NEW: AGENT REVIEW AND APPROVAL =====
@app.post("/agent_review/{lead_id}")
async def agent_review_contract(lead_id: int, action: str = Body(..., embed=True), notes: str = Body("", embed=True)):
    """
    NEW ENDPOINT: Agent reviews generated contract and approves/rejects/requests changes
    Actions: 'approve', 'reject', 'request_changes'
    """
    valid_actions = ['approve', 'reject', 'request_changes']
    if action not in valid_actions:
        raise HTTPException(status_code=400, detail=f"Action must be one of: {valid_actions}")
    
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)
    
    # Determine new status based on action
    status_map = {
        'approve': 'docusign_ready',
        'reject': 'needs_attention',
        'request_changes': 'needs_attention'
    }
    new_status = status_map[action]
    
    async with engine.begin() as conn:
        await conn.execute(
            text("""
                UPDATE leads 
                SET status = :status, agent_notes = :notes, reviewed_at = :reviewed_at
                WHERE id = :id
            """),
            {
                "status": new_status,
                "notes": notes,
                "reviewed_at": datetime.datetime.now(),
                "id": lead_id
            }
        )
    
    await engine.dispose()
    
    return {
        "success": True,
        "lead_id": lead_id,
        "action": action,
        "new_status": new_status,
        "notes": notes
    }

# ===== NEW: DEADLINE MONITORING =====
@app.get("/deadlines")
async def get_upcoming_deadlines():
    """
    NEW ENDPOINT: Get all upcoming deadlines across active transactions
    For agent dashboard and follow-up automation
    """
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.connect() as conn:
        result = await conn.execute(
            text("""
                SELECT id, name, status, created_at, realist_data 
                FROM leads 
                WHERE status IN ('contract_drafted', 'docusign_ready', 'pending_signatures')
            """)
        )
        active_leads = result.fetchall()

    await engine.dispose()

    deadline_summary = []
    today = datetime.date.today()
    
    for lead in active_leads:
        # Calculate deadlines based on contract creation date
        contract_date = lead[3].date() if lead[3] else today
        deadlines = calculate_transaction_deadlines(contract_date)
        
        # Check which deadlines are approaching (within 3 days)
        approaching = []
        for deadline_type, deadline_date in deadlines.items():
            days_until = (deadline_date - today).days
            if 0 <= days_until <= 3:
                approaching.append({
                    "type": deadline_type,
                    "date": deadline_date.isoformat(),
                    "days_until": days_until
                })
        
        if approaching:
            deadline_summary.append({
                "lead_id": lead[0],
                "lead_name": lead[1],
                "status": lead[2],
                "approaching_deadlines": approaching
            })
    
    return {
        "upcoming_deadlines": deadline_summary,
        "total_active_transactions": len(active_leads)
    }

# Enhanced status update endpoint
@app.post("/update_status/{lead_id}")
async def update_status(lead_id: int, new_status: str = Body(..., embed=True)):
    """
    ENHANCED: Update lead status with validation for new workflow states
    """
    valid_statuses = [
        "new", "realist_added", "contract_drafted", "awaiting_review", 
        "docusign_ready", "pending_signatures", "completed", "needs_attention"
    ]
    
    if new_status.lower() not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Status must be one of: {valid_statuses}")
    
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.begin() as conn:
        await conn.execute(
            text("UPDATE leads SET status = :status WHERE id = :id"),
            {"status": new_status.lower(), "id": lead_id}
        )

    await engine.dispose()
    return {"success": True, "updated_id": lead_id, "new_status": new_status}

# ===== NEW: BATCH FOLLOW-UP TRIGGER =====
@app.post("/trigger_followups")
async def trigger_followups():
    """
    NEW ENDPOINT: Manually trigger follow-up sequence for pending items
    Later this will be automated via cron job
    """
    DATABASE_URL = get_database_url()
    engine = create_async_engine(DATABASE_URL, echo=False)

    async with engine.connect() as conn:
        result = await conn.execute(
            text("""
                SELECT id, name, email, status 
                FROM leads 
                WHERE status IN ('pending_signatures', 'awaiting_review')
            """)
        )
        pending_leads = result.fetchall()

    await engine.dispose()

    # TODO: Implement actual email/SMS follow-up logic
    follow_up_results = []
    for lead in pending_leads:
        # Placeholder for follow-up logic
        follow_up_results.append({
            "lead_id": lead[0],
            "name": lead[1],
            "email": lead[2],
            "status": lead[3],
            "action": "follow_up_scheduled"  # Placeholder
        })
    
    return {
        "success": True,
        "follow_ups_triggered": len(follow_up_results),
        "results": follow_up_results
    }

# ===== NEW: HEALTH CHECK FOR INTEGRATIONS =====
@app.get("/health")
async def health_check():
    """
    NEW ENDPOINT: Check system health and integration status
    """
    health_status = {
        "database": "unknown",
        "contract_templates": "unknown",
        "generated_contracts_dir": "unknown"
    }
    
    # Check database
    try:
        DATABASE_URL = get_database_url()
        engine = create_async_engine(DATABASE_URL, echo=False)
        async with engine.connect() as conn:
            await conn.execute(text("SELECT 1"))
        await engine.dispose()
        health_status["database"] = "healthy"
    except:
        health_status["database"] = "error"
    
    # Check contract templates
    template_files = [
        "templates/contracts/Standard_Purchase_Agreement.pdf",
        "templates/contracts/Exclusive Right to Represent Buyer Brokerage Agreement (BBA-SA) .pdf",
        "templates/contracts/Exclusive Right to Sell Brokerage Agreement to Standard Listing%.pdf"
    ]
    
    available_templates = [f for f in template_files if os.path.exists(f)]
    health_status["contract_templates"] = f"{len(available_templates)}/{len(template_files)} available"
    
    # Check generated contracts directory
    health_status["generated_contracts_dir"] = "exists" if os.path.exists("generated_contracts") else "missing"
    
    return {
        "status": "Mira v1.5 Health Check",
        "timestamp": datetime.datetime.now().isoformat(),
        "health": health_status
    }